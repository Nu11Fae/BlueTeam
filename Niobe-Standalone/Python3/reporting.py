from __future__ import annotations

import contextlib
import html
import io
import json
import os
import re
import shutil
import subprocess
import tempfile
import textwrap
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from jinja2 import Environment, FileSystemLoader, select_autoescape
from markdown import markdown as md_render
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from .hld_baseline import COMPLIANCE_STANDARDS, SECURITY_VIEW_STANDARDS
from .reference_docs import load_reference_bundle
from .scoring import assign_grade
from .settings import get_settings

_ACTIVITY_TYPE = os.environ.get("NIOBE_ACTIVITY_TYPE", "Digital-Audit")
TITLE = f"AI Technical Intelligence Review — {_ACTIVITY_TYPE}"
CLASSIFICATION = "CONFIDENZIALE"
OFFICIAL_SECTIONS = {
    "1": "1. EXECUTIVE SUMMARY",
    "2": "2. SCOPE, METODOLOGIA E LIMITAZIONI",
    "3": "3. TOPOLOGIA E HEATMAP DEL RISCHIO",
    "4": "4. INTELLIGENCE ARCHITETTURALE",
    "5": "5. SECURITY VIEW",
    "6": "6. DATI, DIPENDENZE E SUPPLY CHAIN",
    "7": "7. ANALISI OSS (Open Source Scan)",
    "8": "8. MATURITA DEVOPS",
    "9": "9. DEBITO TECNICO E MANUTENIBILITA",
    "10": "10. ANALISI DI COMPLIANCE",
    "11": "11. ANALISI DEI FINDING",
    "12": "12. CONCLUSIONI",
}
OUTLINE_TITLES = [OFFICIAL_SECTIONS[str(index)] for index in range(1, 13)]
ITALIAN_MONTHS = {
    1: "gennaio",
    2: "febbraio",
    3: "marzo",
    4: "aprile",
    5: "maggio",
    6: "giugno",
    7: "luglio",
    8: "agosto",
    9: "settembre",
    10: "ottobre",
    11: "novembre",
    12: "dicembre",
}
PRIMARY_BLUE = "1B365D"
SECONDARY_BLUE = "5F7EA6"
TABLE_BORDER = "A9B8C8"
TITLE_FONT = "Avenir Next LT Pro"
HEADING_FONT = "Aptos"
SUBHEADING_FONT = "Aptos"
BODY_FONT = "Aptos"
CODE_FONT = "Aptos Mono"
PLACEHOLDER_MARKERS = (
    "placeholder",
    "inserire qui",
    "inserire placeholder",
    "sostituire con",
    "nota metodologica",
    "didascalia placeholder",
    "placeholder_tag",
    "bozza da completare",
)


def _slug_anchor(value: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-")
    return slug or "section"


def _outline() -> list[dict[str, str]]:
    return [{"title": title, "anchor": _slug_anchor(title)} for title in OUTLINE_TITLES]


def _read_text(pathish: object | None) -> str:
    if not pathish:
        return ""
    candidate = Path(str(pathish))
    if candidate.exists() and candidate.is_file():
        return candidate.read_text(encoding="utf-8")
    return ""


def _italian_date(value: datetime | None = None) -> str:
    current = value or datetime.now(UTC)
    return f"{current.day:02d} {ITALIAN_MONTHS[current.month]} {current.year}"


def _mask_remediation_dates(markdown: str) -> str:
    date_patterns = [
        re.compile(r"\b\d{4}-\d{2}-\d{2}\b"),
        re.compile(r"\b\d{1,2}/\d{1,2}/\d{2,4}\b"),
    ]
    guarded: list[str] = []
    for line in markdown.splitlines():
        lowered = line.lower()
        if any(token in lowered for token in ("remediation", "mitigation", "fix", "timeline", "scadenza", "remediation window")):
            for pattern in date_patterns:
                line = pattern.sub("finestra di remediation definita", line)
        guarded.append(line)
    return "\n".join(guarded)


def _split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _format_table_row(cells: list[str]) -> str:
    return "| " + " | ".join(cells) + " |"


def _strip_effort_columns(markdown: str) -> str:
    lines = markdown.splitlines()
    rewritten: list[str] = []
    index = 0
    while index < len(lines):
        stripped = lines[index].strip()
        if stripped.startswith("|"):
            block: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                block.append(lines[index])
                index += 1
            if len(block) >= 2:
                headers = _split_table_row(block[0])
                separator = _split_table_row(block[1])
                if headers and separator and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in separator):
                    keep = [
                        idx for idx, header in enumerate(headers)
                        if not any(token in header.lower() for token in ("effort", "gg/p", "giornat", "man/day", "day effort"))
                    ]
                    if keep and len(keep) != len(headers):
                        rewritten.append(_format_table_row([headers[idx] for idx in keep]))
                        rewritten.append(_format_table_row(["---" for _ in keep]))
                        for row in block[2:]:
                            cells = _split_table_row(row)
                            cells += [""] * (len(headers) - len(cells))
                            rewritten.append(_format_table_row([cells[idx] for idx in keep]))
                        continue
            rewritten.extend(block)
            continue
        rewritten.append(lines[index])
        index += 1
    return "\n".join(rewritten)


def _normalize_heading(title: str) -> str:
    match = re.match(r"^(\d+)\.\s+", title.strip())
    if not match:
        return title.strip()
    return OFFICIAL_SECTIONS.get(match.group(1), title.strip())


def _sanitize_llm_markdown(markdown: str) -> str:
    if not markdown.strip():
        return ""
    cleaned: list[str] = []
    skip_level = 0
    for raw_line in html.unescape(markdown).splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            if cleaned and cleaned[-1] != "":
                cleaned.append("")
            continue
        lowered = stripped.lower()
        if any(marker in lowered for marker in PLACEHOLDER_MARKERS):
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            title = heading.group(2).strip()
            lowered = title.lower()
            if skip_level and level <= skip_level:
                skip_level = 0
            if lowered == TITLE.lower():
                continue
            if lowered in {"indice", "index", "appendici", "13. appendici"}:
                skip_level = level
                continue
            if "vibecoding" in lowered or "vibe coding" in lowered or "rapid review" in lowered or "full review" in lowered:
                skip_level = level
                continue
            if lowered.startswith("11. prospettive"):
                skip_level = level
                continue
            if lowered.startswith("12. "):
                line = f"{heading.group(1)} {OFFICIAL_SECTIONS['11']}"
            elif re.match(r"^\d+\.\s+", title):
                line = f"{heading.group(1)} {_normalize_heading(title)}"
            elif lowered == "security risk assessment (owasp top 10)":
                line = f"{heading.group(1)} {OFFICIAL_SECTIONS['5']}"
        if skip_level:
            continue
        cleaned.append(line)
    body = "\n".join(cleaned).strip()
    body = re.sub(r"\n{3,}", "\n\n", body)
    body = body.replace("## 10. RISK FINDINGS PRIORITARI", f"## {OFFICIAL_SECTIONS['11']}")
    body = body.replace("## 11. RISK FINDINGS PRIORITARI", f"## {OFFICIAL_SECTIONS['11']}")
    body = re.sub(r"^##\s+11\.\s+.*FINDING.*$", f"## {OFFICIAL_SECTIONS['11']}", body, flags=re.M | re.I)
    body = _strip_effort_columns(body)
    return _mask_remediation_dates(body)


def _markdown_to_html(markdown_text: str) -> str:
    return md_render(markdown_text, extensions=["tables", "fenced_code", "sane_lists", "nl2br"], output_format="html5")


def _extract_sections(markdown: str) -> list[dict[str, str]]:
    sections: list[dict[str, str]] = []
    current: dict[str, Any] | None = None
    for raw_line in markdown.splitlines():
        stripped = raw_line.strip()
        heading = re.match(r"^##\s+(.*)$", stripped)
        if heading and re.match(r"^\d+\.\s+", heading.group(1).strip()):
            title = _normalize_heading(heading.group(1).strip())
            number = title.split(".", 1)[0]
            if number not in OFFICIAL_SECTIONS:
                continue
            if current is not None:
                body = "\n".join(current["body_lines"]).strip()
                current["body_markdown"] = body
                current["html"] = _markdown_to_html(body) if body else ""
                current.pop("body_lines", None)
                sections.append(current)  # type: ignore[arg-type]
            current = {
                "title": title,
                "number": number,
                "anchor": _slug_anchor(title),
                "body_lines": [],
            }
            continue
        if current is not None:
            current["body_lines"].append(raw_line)
    if current is not None:
        body = "\n".join(current["body_lines"]).strip()
        current["body_markdown"] = body
        current["html"] = _markdown_to_html(body) if body else ""
        current.pop("body_lines", None)
        sections.append(current)  # type: ignore[arg-type]
    return sections


def extract_report_sections(markdown: str) -> list[dict[str, str]]:
    return _extract_sections(markdown)


def _tool_status_counts(tool_results: dict[str, object]) -> tuple[int, int, int]:
    completed = failed = skipped = 0
    for raw in tool_results.values():
        result = raw if isinstance(raw, dict) else {}
        if result.get("skipped"):
            skipped += 1
        elif result.get("ok"):
            completed += 1
        else:
            failed += 1
    return completed, failed, skipped


def _escape_svg(text: str) -> str:
    return html.escape(text, quote=True)


def _label_lines(text: str, width: int = 18) -> list[str]:
    words = text.split()
    if not words:
        return [""]
    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        candidate = f"{current} {word}"
        if len(candidate) <= width:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines[:3]


def _svg_text_block(x: int, y: int, text: str, line_gap: int = 18, cls: str = "") -> str:
    lines = _label_lines(text)
    spans = []
    for index, line in enumerate(lines):
        dy = "0" if index == 0 else str(line_gap)
        spans.append(f'<tspan x="{x}" dy="{dy}">{_escape_svg(line)}</tspan>')
    class_attr = f' class="{cls}"' if cls else ""
    return f'<text text-anchor="middle" x="{x}" y="{y}"{class_attr}>' + "".join(spans) + "</text>"


def _application_type(repo_profile: dict[str, Any]) -> str:
    languages = repo_profile.get("languages", []) or []
    frameworks = repo_profile.get("frameworks", []) or []
    if ".NET" in languages:
        return "Applicativo .NET"
    if "Python" in languages and "JavaScript/TypeScript" in languages:
        return "Applicativo full-stack"
    if "Python" in languages:
        return "Applicativo Python"
    if "JavaScript/TypeScript" in languages:
        return "Applicativo Node.js / front-end"
    if "Rust" in languages:
        return "Applicativo Rust / servizi nativi"
    if "C/C++" in languages:
        return "Applicativo nativo C/C++"
    if frameworks:
        return f"Applicativo {frameworks[0]}"
    return "Applicativo software custom"


def _environment_summary(context: dict[str, Any]) -> str:
    manifest = context.get("manifest", {}) if isinstance(context.get("manifest", {}), dict) else {}
    repo_profile = context.get("repo_profile", {}) if isinstance(context.get("repo_profile", {}), dict) else {}
    codebase_metrics = context.get("codebase_metrics", {}) if isinstance(context.get("codebase_metrics", {}), dict) else {}
    host_env = manifest.get("host_environment", {}) if isinstance(manifest.get("host_environment", {}), dict) else {}
    system = str(host_env.get("system") or "Host non dichiarato")
    shell = str(host_env.get("shell") or "shell non dichiarata")
    languages = ", ".join(repo_profile.get("languages", []) or []) or "stack non determinato"
    loc = codebase_metrics.get("code_lines", 0)
    app_type = _application_type(repo_profile)
    return f"{system} | Digital Audit containerizzato | ~{loc} LOC core | {languages} | {app_type} | shell {shell}"


def _recipient_label(context: dict[str, Any]) -> str:
    reference_documents = context.get("reference_documents", {}) if isinstance(context.get("reference_documents", {}), dict) else {}
    engagement = reference_documents.get("engagement", {}) if isinstance(reference_documents.get("engagement", {}), dict) else {}
    if engagement.get("recipients"):
        return str(engagement["recipients"])
    client_name = str(context.get("client_name", "Cliente"))
    first_name = str(engagement.get("contact_first_name", "")).strip()
    last_name = str(engagement.get("contact_last_name", "")).strip()
    if first_name or last_name:
        return f"Cliente {client_name}, referente {first_name} {last_name}".strip()
    return f"Cliente {client_name}"


def _overall_rating(context: dict[str, Any]) -> dict[str, Any]:
    findings = context.get("risk_findings", []) if isinstance(context.get("risk_findings", []), list) else []
    if not findings:
        return {"grade": "N/D", "score": 0.0, "validated": 0, "candidate": 0}
    scored = [float(item.get("final_score", 0.0)) for item in findings if isinstance(item, dict)]
    score = max(scored) if scored else 0.0
    grade = assign_grade(score, get_settings()) if score else "N/D"
    validated = sum(1 for item in findings if isinstance(item, dict) and item.get("validation_status") == "Validated")
    candidate = sum(1 for item in findings if isinstance(item, dict) and item.get("validation_status") == "Candidate")
    return {"grade": grade, "score": round(score, 2), "validated": validated, "candidate": candidate}


def _delivery_inventory_rows() -> list[dict[str, str]]:
    return [
        {
            "artifact": "AI Technical Intelligence Review.docx",
            "description": "Documento finale Word generato sul template canonico Tecnolife e pronto per revisione o stampa.",
            "verification": 'gpg --auto-key-import --verify signatures/AI Technical Intelligence Review.docx.asc "AI Technical Intelligence Review.docx"',
        },
        {
            "artifact": "AI Technical Intelligence Review.pdf",
            "description": "Documento finale PDF derivato dal DOCX canonico e firmato digitalmente.",
            "verification": 'gpg --auto-key-import --verify signatures/AI Technical Intelligence Review.pdf.asc "AI Technical Intelligence Review.pdf"',
        },
        {
            "artifact": "Risk Register.xlsx",
            "description": "Registro dei rischi con scoring HLD, grade A-E e classificazione transazionale.",
            "verification": 'gpg --auto-key-import --verify signatures/Risk Register.xlsx.asc "Risk Register.xlsx"',
        },
        {
            "artifact": "sbom.cyclonedx.json",
            "description": "SBOM machine-readable in formato CycloneDX.",
            "verification": "gpg --auto-key-import --verify signatures/sbom.cyclonedx.json.asc sbom.cyclonedx.json",
        },
        {
            "artifact": "Compliance Artifact.pdf",
            "description": "Estratto documentale dedicato all’analisi di compliance.",
            "verification": 'gpg --auto-key-import --verify signatures/Compliance Artifact.pdf.asc "Compliance Artifact.pdf"',
        },
        {
            "artifact": "OSS Provenance Report (Copyleft Risk).pdf",
            "description": "Sintesi documentale della provenance OSS e dei vincoli licenza.",
            "verification": 'gpg --auto-key-import --verify signatures/OSS Provenance Report (Copyleft Risk).pdf.asc "OSS Provenance Report (Copyleft Risk).pdf"',
        },
        {
            "artifact": "evidence-manifest.json",
            "description": "Manifest forense della sessione con hash, identificativi e metadati di esecuzione.",
            "verification": "gpg --auto-key-import --verify signatures/evidence-manifest.json.asc evidence-manifest.json",
        },
        {
            "artifact": "signatures/",
            "description": "Firme detached GPG e marche temporali OTS dei deliverable top-level.",
            "verification": "gpg --auto-key-import --verify signatures/<file>.asc <file> && ots verify signatures/<file>.asc.ots",
        },
        {
            "artifact": "AI_DA.zip / AI_DA.zip.asc",
            "description": "Pacchetto finale ricorsivo della delivery e relativa firma detached con include-key-block.",
            "verification": "gpg --auto-key-import --verify AI_DA.zip.asc AI_DA.zip",
        },
    ]


def _legend_rows() -> list[dict[str, str]]:
    return [
        {
            "term": "Evidence",
            "definition": "Elemento documentale o tecnico acquisito durante il workflow di audit e correlato a un artefatto verificabile.",
        },
        {
            "term": "Candidate Finding",
            "definition": "Per Candidate Finding si intende una possibile vulnerabilità, carenza di controllo, debolezza architetturale, esposizione di compliance o anomalia operativa emersa da attività di AI-assisted review, code review, DAST, analisi documentale, review architetturale o validazione analitica. Il Candidate Finding ha natura istruttoria e non è, di per sé, sufficiente per l’iscrizione nel Risk Register.",
        },
        {
            "term": "Validated Finding",
            "definition": "Il Candidate Finding diviene Validated Finding solo quando risulta almeno localizzato su asset, componente, flusso o controllo specifico; supportato da evidenza sufficiente o da evidence summary verificabile; classificato secondo una tassonomia tecnica o di controllo, ove disponibile; accompagnato da una spiegazione del possibile impatto tecnico, business e, se del caso, regolatorio; corredato da un livello di confidence espressamente dichiarato.",
        },
        {
            "term": "Likelihood / Exploitability",
            "definition": "Probabilità ragionevole di sfruttamento, abuso, attivazione o manifestazione della debolezza, tenuto conto di exploitability, esponibilità, semplicità dell’attacco, prerequisiti e realisticità dello scenario.",
        },
        {
            "term": "Technical Impact",
            "definition": "Impatto tecnico potenziale su confidenzialità, integrità, disponibilità, autenticazione, autorizzazione, segregazione, tracciabilità, affidabilità applicativa o resilienza del sistema.",
        },
        {
            "term": "Business Impact",
            "definition": "Rilevanza dell’asset, del componente, del processo o del flusso colpito rispetto a revenue generation, operations, clienti, SLA, dati critici, servizi core o processi core business.",
        },
        {
            "term": "Control Weakness",
            "definition": "Livello di debolezza dei controlli esistenti, tenendo conto di controlli preventivi, detective, compensativi, segregazione, logging, monitoraggio, design difensivo e capacità reale di riduzione del rischio.",
        },
        {
            "term": "Compliance Exposure",
            "definition": "Possibile esposizione a obblighi normativi, regolatori, contrattuali o standard di controllo applicabili al target, al servizio o al cliente.",
        },
        {
            "term": "Remediation Effort",
            "definition": "Sforzo atteso e ipotizzato di remediation in termini di costo, tempo, complessità tecnica, dipendenze architetturali, refactoring, redesign, re-test, migration effort o change impact.",
        },
        {
            "term": "Transaction Impact",
            "definition": "Possibile incidenza del finding sulla transazione, sulla negoziazione, sulla valuation, sul piano di remediation, sul sign-off, sul contratto stesso del cliente per l’acquisizione.",
        },
        {
            "term": "Evidence Confidence",
            "definition": "L’Evidence Confidence è espressa separatamente come High / Medium / Low e non viene trattata come sinonimo di rischio. Essa misura la robustezza probatoria del finding e non la sua materialità.",
        },
        {
            "term": "Red Flag",
            "definition": "Finding classificato almeno come D o E, o come Red Flag autonoma, quando presenta caratteri tali da incidere materialmente su signing, closing, prezzo, protezioni contrattuali, continuità del servizio o affidabilità sostanziale del prodotto.",
        },
    ]


def _rating_dimension_rows() -> list[dict[str, str]]:
    return [
        {"label": "Likelihood / Exploitability", "definition": "Probabilità ragionevole di sfruttamento, abuso, attivazione o manifestazione della debolezza, tenuto conto di exploitability, esponibilità, semplicità dell’attacco, prerequisiti e realisticità dello scenario."},
        {"label": "Technical Impact", "definition": "Impatto tecnico potenziale su confidenzialità, integrità, disponibilità, autenticazione, autorizzazione, segregazione, tracciabilità, affidabilità applicativa o resilienza del sistema."},
        {"label": "Business Impact", "definition": "Rilevanza dell’asset, del componente, del processo o del flusso colpito rispetto a revenue generation, operations, clienti, SLA, dati critici, servizi core o processi core business."},
        {"label": "Control Weakness", "definition": "Livello di debolezza dei controlli esistenti, tenendo conto di controlli preventivi, detective, compensativi, segregazione, logging, monitoraggio, design difensivo e capacità reale di riduzione del rischio."},
        {"label": "Compliance Exposure", "definition": "Possibile esposizione a obblighi normativi, regolatori, contrattuali o standard di controllo applicabili al target, al servizio o al cliente."},
        {"label": "Remediation Effort", "definition": "Sforzo atteso e ipotizzato di remediation in termini di costo, tempo, complessità tecnica, dipendenze architetturali, refactoring, redesign, re-test, migration effort o change impact."},
        {"label": "Transaction Impact", "definition": "Possibile incidenza del finding sulla transazione, sulla negoziazione, sulla valuation, sul piano di remediation, sul sign-off, sul contratto stesso del cliente per l’acquisizione."},
        {"label": "Evidence Confidence", "definition": "Livello di confidenza della conclusione, espresso separatamente come High / Medium / Low, senza confondere incertezza probatoria e severità del rischio."},
    ]


def _rating_logic_rows() -> list[dict[str, str]]:
    return [
        {"label": "Inherent Risk", "formula": "0,35 × Likelihood + 0,40 × Technical Impact + 0,25 × Business Impact"},
        {"label": "Residual Risk Score", "formula": "0,70 × Inherent Risk + 0,30 × Control Weakness"},
        {"label": "Transaction Materiality", "formula": "0,50 × Transaction Impact + 0,25 × Compliance Exposure + 0,25 × Remediation Effort"},
        {"label": "Final Score", "formula": "MAX(Residual Risk Score, Transaction Materiality)"},
        {"label": "Arrotondamento", "formula": "I risultati intermedi restano a due decimali; l’arrotondamento alla prima cifra decimale si applica solo per la mappatura in classe A–E."},
    ]


def _grade_meaning_rows() -> list[dict[str, str]]:
    return [
        {"grade": "A", "meaning": "Acceptable / Non rilevante: rischio trascurabile, localizzato, con impatto minimo o nullo, controlli efficaci e remediation ordinaria."},
        {"grade": "B", "meaning": "Low: rischio limitato, non sistemico, senza impatto materiale sul deal, con remediation rapida e costo contenuto."},
        {"grade": "C", "meaning": "Moderate: rischio rilevante ma gestibile, con impatto tecnico o regolatorio non trascurabile, remediation non banale e possibile impatto su piano di integrazione o costo operativo."},
        {"grade": "D", "meaning": "High: rischio materiale per il cliente, per asset o processi critici o per la compliance, con remediation onerosa o lenta e potenziale riflesso su negoziazione, valuation, covenants o remediation pre-close."},
        {"grade": "E", "meaning": "Critical / Deal-Significant: rischio critico, suscettibile di incidere su signing, closing, prezzo, protezioni contrattuali, continuità del servizio o affidabilità sostanziale del prodotto."},
    ]


def _rating_logic_intro() -> str:
    return (
        "Il modello HLD non utilizza una sola matrice binaria Probabilità × Impatto, ma una struttura composita che distingue Inherent Risk, Residual Risk Score e Transaction Materiality. "
        "La finalità è separare il rischio intrinseco del finding, la qualità dei controlli esistenti e la sua materialità economico-transazionale, mantenendo il calcolo verificabile ex post e replicabile da un reviewer indipendente."
    )


def _rating_logic_clauses() -> str:
    return (
        "Il modello opera sul piano numerico come algoritmo deterministico di scoring ponderato, applicato a valori discreti da 1 a 5. "
        "L’ordine logico obbligato è il seguente: (i) calcolo dell’Inherent Risk come media ponderata di Likelihood, Technical Impact e Business Impact; "
        "(ii) calcolo del Residual Risk incorporando la Control Weakness; (iii) calcolo separato della Transaction Materiality sulla base di Transaction Impact, Compliance Exposure e Remediation Effort; "
        "(iv) definizione del Final Score come valore maggiore fra Residual Risk e Transaction Materiality. "
        "Salvo diversa espressa indicazione, i risultati intermedi sono mantenuti con due decimali ai fini di tracciabilità del calcolo; l’eventuale arrotondamento alla prima cifra decimale viene applicato solo ai fini della mappatura in classe A–E."
    )


def _disclaimer_text() -> str:
    return (
        "La presente Liberatoria deve essere letta insieme alla definizione di Evidence riportata nell’HLD e al manifesto delle evidenze allegato. "
        "Il presente Digital Audit è una valutazione professionale evidence-driven, costruita su standard tecnici e normativi riconosciuti, su una pipeline di acquisizione e correlazione delle evidenze e su controlli di firma e marcatura temporale. "
        "L’analisi non costituisce certificazione di assenza assoluta di vulnerabilità, difetti logici, anomalie operative o non conformità al di fuori del perimetro esaminato, né sostituisce penetration test estensivi, assessment infrastrutturali o verifiche runtime full-scope.\n\n"
        "Le conclusioni dipendono dalla qualità e completezza del repository, degli artefatti acquisiti, delle configurazioni rese disponibili e delle evidenze effettivamente raccolte durante la finestra di audit. "
        "Eventuali percorsi non raggiungibili staticamente, comportamenti runtime non documentati, configurazioni esterne non consegnate o integrazioni non osservabili possono ridurre la completezza della valutazione.\n\n"
        "Tecnolife ha adottato procedure ragionevoli e verificabili per preservare l’affidabilità delle evidenze; resta tuttavia esclusa ogni garanzia di esaustività assoluta su elementi non inclusi nel perimetro concordato o non osservabili con le regole di ingaggio applicate."
    )


def _evidence_manifest_intro() -> str:
    return (
        "L’Evidence Manifest è il documento strutturato che cristallizza gli elementi forensi della sessione di audit: identificativi host e run-time, timestamp, hash del repository, metadati di acquisizione e correlazione tra perimetro analizzato e artefatti prodotti. "
        "La sua funzione è consentire a un reviewer terzo di ricostruire in modo verificabile quando, dove e su quale contenuto è stata eseguita l’analisi."
    )


def _finding_definition_block() -> str:
    return textwrap.dedent(
        """
        ### Definizione del finding

        Per Candidate Finding si intende una possibile vulnerabilità, carenza di controllo, debolezza architetturale, esposizione di compliance o anomalia operativa emersa da attività di AI-assisted review, code review, DAST, analisi documentale, review architetturale o validazione analitica. Il Candidate Finding ha natura istruttoria e non è, di per sé, sufficiente per l’iscrizione nel Risk Register.

        Il Candidate Finding diviene Validated Finding solo quando risulta almeno:
        - localizzato su asset, componente, flusso o controllo specifico;
        - supportato da evidenza sufficiente o da evidence summary verificabile;
        - classificato secondo una tassonomia tecnica o di controllo, ove disponibile;
        - accompagnato da una spiegazione del possibile impatto tecnico, business e, se del caso, regolatorio;
        - corredato da un livello di confidence espressamente dichiarato.
        """
    ).strip()


def _fallback_body(context: dict[str, Any]) -> str:
    summary = context.get("summary", {}) if isinstance(context.get("summary", {}), dict) else {}
    repo_profile = context.get("repo_profile", {}) if isinstance(context.get("repo_profile", {}), dict) else {}
    codebase_metrics = context.get("codebase_metrics", {}) if isinstance(context.get("codebase_metrics", {}), dict) else {}
    languages = ", ".join(repo_profile.get("languages", []) or []) or "stack non determinato"
    return textwrap.dedent(
        f"""
        ## 1. EXECUTIVE SUMMARY
        ### Sintesi per il management
        {summary.get('executive_summary', 'Il Digital Audit ha completato acquisizione forense delle evidenze, correlazione dei finding e revisione LLM evidence-driven sul perimetro applicativo.')}

        ### Dashboard esecutiva
        - Linguaggi rilevati: {languages}
        - LOC core: {codebase_metrics.get('code_lines', 0)}
        - Strato di analisi: Digital Audit containerizzato con revisione LLM post-correlazione

        ## 2. SCOPE, METODOLOGIA E LIMITAZIONI
        ### 2.1 Valore probatorio dell’analisi
        L’attività si fonda su hashing SHA-512 del perimetro, manifest di esecuzione, identificativi host e firme detached con marca temporale.

        ### 2.4 Core Application Logic
        Il focus ricade sui componenti custom applicativi, distinguendoli da bootstrap, generated code, vendor code e materiale legacy.

        ## 3. TOPOLOGIA E HEATMAP DEL RISCHIO
        La topologia e la heatmap sono derivate da path core, dipendenze, finding clusterizzati e segnali statici correlati.

        ## 4. INTELLIGENCE ARCHITETTURALE
        L’analisi architetturale distingue debolezze strutturali, debt accumulato e limiti di scalabilità o governabilità.

        ## 5. SECURITY VIEW
        La Security View è costruita su evidenze tool-backed, revisione profonda del codice e standard HLD.

        ## 6. DATI, DIPENDENZE E SUPPLY CHAIN
        La supply chain è valutata tramite SBOM, provenance OSS, segnali licensing e possibili blocker transazionali.

        ## 7. ANALISI OSS (Open Source Scan)
        L’analisi OSS è costruita sulle evidenze di ScanCode, FOSSology, SBOM CycloneDX per identificare licenze restrittive e snippet non dichiarati.

        ## 8. MATURITA DEVOPS
        La maturità DevOps è desunta dalle evidenze presenti nel repository e dalle pratiche rilevabili staticamente.

        ## 9. DEBITO TECNICO E MANUTENIBILITA
        Il debito tecnico osservato si riflette nella distribuzione dei path core, nella densità di finding e nei segnali statici di maintainability.

        ## 10. ANALISI DI COMPLIANCE
        L’analisi di compliance correla codice, dipendenze, dati e scelte architetturali rispetto ai framework HLD applicabili.

        ## 11. ANALISI DEI FINDING
        I finding prioritari sono raccolti nel Risk Register allegato e devono essere letti con la relativa classificazione transazionale.

        ## 12. CONCLUSIONI
        Le conclusioni devono essere lette insieme alla delivery firmata e alle appendici metodologiche del presente report.
        """
    ).strip()


def _scope_label(context: dict[str, Any]) -> str:
    return _environment_summary(context)


def _topology_svg(context: dict[str, Any]) -> str:
    repo_profile = context.get("repo_profile", {}) if isinstance(context.get("repo_profile", {}), dict) else {}
    languages = repo_profile.get("languages", []) or ["Stack non determinato"]
    frameworks = repo_profile.get("frameworks", []) or ["Nessun orchestratore dichiarato"]
    notes = repo_profile.get("notes", []) or ["Nessuna nota addizionale sul perimetro core."]
    source_paths = repo_profile.get("source_paths", []) or [str(context.get("project_name", "target"))]
    exclusions = [Path(path).name for path in (repo_profile.get("candidate_exclusions", []) or [])] or ["Nessuna esclusione candidata"]
    codebase_metrics = context.get("codebase_metrics", {}) if isinstance(context.get("codebase_metrics", {}), dict) else {}
    core_node = f"{Path(str(context.get('analysis_target_path') or context.get('project_name', 'Target'))).name} core"
    evidence_node = f"{codebase_metrics.get('files', 0)} file / {codebase_metrics.get('code_lines', 0)} LOC"
    return f'''<svg class="diagram-svg" viewBox="0 0 900 430" role="img" aria-label="Topologia del rischio derivata da evidenze repository-centriche">
  <defs>
    <linearGradient id="tl-grad" x1="0" y1="0" x2="1" y2="1">
      <stop offset="0%" stop-color="#eef4fb" />
      <stop offset="100%" stop-color="#dbe7f7" />
    </linearGradient>
    <marker id="tl-arrow" markerWidth="10" markerHeight="10" refX="8" refY="5" orient="auto">
      <path d="M0,0 L10,5 L0,10 z" fill="#5d6b80"></path>
    </marker>
  </defs>
  <rect x="42" y="150" width="170" height="70" rx="16" fill="url(#tl-grad)" stroke="#b6c7dd" />
  {_svg_text_block(127, 182, str(context.get('client_name', 'Stakeholder')), cls='node-text')}
  <rect x="330" y="118" width="250" height="120" rx="20" fill="#1b365d" />
  {_svg_text_block(455, 165, core_node, cls='center-text')}
  {_svg_text_block(455, 205, ' / '.join(Path(path).name for path in source_paths[:3]), 16, 'center-subtext')}
  <rect x="675" y="45" width="175" height="68" rx="16" fill="url(#tl-grad)" stroke="#b6c7dd" />
  {_svg_text_block(762, 75, 'Linguaggi: ' + ', '.join(languages[:3]), cls='node-text')}
  <rect x="675" y="132" width="175" height="68" rx="16" fill="url(#tl-grad)" stroke="#b6c7dd" />
  {_svg_text_block(762, 162, 'Framework: ' + ', '.join(frameworks[:3]), cls='node-text')}
  <rect x="675" y="219" width="175" height="68" rx="16" fill="url(#tl-grad)" stroke="#b6c7dd" />
  {_svg_text_block(762, 249, 'Evidence pack: ' + evidence_node, cls='node-text')}
  <rect x="675" y="306" width="175" height="82" rx="16" fill="#fdf4f2" stroke="#e6b8b2" />
  {_svg_text_block(762, 337, 'Esclusioni: ' + ', '.join(exclusions[:3]), cls='warn-text')}
  <path d="M212 185 L330 178" stroke="#5d6b80" stroke-width="3" fill="none" marker-end="url(#tl-arrow)" />
  <path d="M580 158 C620 130 635 108 675 82" stroke="#5d6b80" stroke-width="3" fill="none" marker-end="url(#tl-arrow)" />
  <path d="M580 174 L675 166" stroke="#8a97aa" stroke-width="3" fill="none" marker-end="url(#tl-arrow)" />
  <path d="M580 190 C624 220 640 236 675 252" stroke="#5d6b80" stroke-width="3" fill="none" marker-end="url(#tl-arrow)" />
  <path d="M560 226 C608 272 630 302 675 343" stroke="#9a6470" stroke-width="3" fill="none" marker-end="url(#tl-arrow)" />
  <rect x="112" y="287" width="470" height="88" rx="18" fill="#eef4fb" stroke="#c7d3e5" />
  {_svg_text_block(347, 322, notes[0], 16, 'note-text')}
</svg>'''


def _heatmap_points(context: dict[str, Any]) -> list[dict[str, Any]]:
    repo_profile = context.get("repo_profile", {}) if isinstance(context.get("repo_profile", {}), dict) else {}
    tool_results = context.get("tool_results", {}) if isinstance(context.get("tool_results", {}), dict) else {}
    codebase_metrics = context.get("codebase_metrics", {}) if isinstance(context.get("codebase_metrics", {}), dict) else {}
    exclusions = repo_profile.get("candidate_exclusions", []) or []
    scale = min(max(int(codebase_metrics.get("code_lines", 0) / 2000), 0), 6)
    app_ok = bool((tool_results.get("semgrep-code") or {}).get("ok")) and bool((tool_results.get("sonarqube") or {}).get("ok"))
    supply_ok = bool((tool_results.get("syft-cyclonedx") or {}).get("artifact_ready")) and bool((tool_results.get("scancode") or {}).get("artifact_ready"))
    return [
        {"label": "Superficie applicativa", "x": 675 if app_ok else 590, "y": 185 if app_ok else 245},
        {"label": "Supply chain", "x": 620 if supply_ok else 545, "y": 255 if supply_ok else 315},
        {"label": "Perimetro legacy", "x": 520 if exclusions else 360, "y": 210 if exclusions else 430},
        {"label": "Debito scalabilita", "x": 580 + scale * 18, "y": 325 - scale * 10},
        {"label": "Compliance", "x": 600 if supply_ok else 510, "y": 300 if exclusions else 360},
    ]


def _heatmap_svg(context: dict[str, Any]) -> str:
    markers = []
    for point in _heatmap_points(context):
        x = int(point["x"])
        y = int(point["y"])
        label = _escape_svg(str(point["label"]))
        markers.append(f'<circle cx="{x}" cy="{y}" r="10" fill="#1b365d"></circle><text x="{x - 110}" y="{y - 14}" class="marker-label">{label}</text>')
    return f'''<svg class="heatmap-svg" viewBox="0 0 900 700" role="img" aria-label="Heatmap del rischio per componenti">
  <defs><linearGradient id="hm-soft" x1="0" y1="0" x2="1" y2="1"><stop offset="0%" stop-color="#fbfcfe" /><stop offset="100%" stop-color="#f1f5fb" /></linearGradient></defs>
  <rect x="70" y="82" width="760" height="520" fill="url(#hm-soft)" stroke="#b8c7dc" stroke-width="2"></rect>
  <line x1="450" y1="82" x2="450" y2="602" stroke="#c0cede" stroke-width="2"></line>
  <line x1="70" y1="342" x2="830" y2="342" stroke="#c0cede" stroke-width="2"></line>
  <text x="450" y="40" text-anchor="middle" class="heat-title">Heatmap Rischi per componenti</text>
  <text x="230" y="130" text-anchor="middle" class="quad-label">Monitorare</text>
  <text x="645" y="130" text-anchor="middle" class="quad-label">Critici da presidiare</text>
  <text x="230" y="382" text-anchor="middle" class="quad-label">Accettabile</text>
  <text x="645" y="382" text-anchor="middle" class="quad-label">Quick win</text>
  <text x="34" y="322" transform="rotate(-90 34 322)" class="axis-label">Impatto</text>
  <text x="450" y="660" text-anchor="middle" class="axis-label">Probabilita</text>
  <text x="178" y="640" text-anchor="middle" class="axis-note">bassa</text>
  <text x="730" y="640" text-anchor="middle" class="axis-note">alta</text>
  {''.join(markers)}
</svg>'''


def _language_svg(context: dict[str, Any]) -> str:
    codebase_metrics = context.get("codebase_metrics", {}) if isinstance(context.get("codebase_metrics", {}), dict) else {}
    by_language = codebase_metrics.get("by_language", {}) or {}
    if not isinstance(by_language, dict) or not by_language:
        return '<svg class="bars-svg" viewBox="0 0 900 280" role="img" aria-label="Distribuzione linguaggi non disponibile"><text x="450" y="140" text-anchor="middle" class="axis-label">Distribuzione linguaggi non disponibile nel pacchetto evidenze.</text></svg>'
    items = sorted(((language, counts.get("code_lines", 0)) for language, counts in by_language.items() if isinstance(counts, dict)), key=lambda item: item[1], reverse=True)[:6]
    max_value = max((value for _, value in items), default=1)
    bars = []
    y = 58
    palette = ["#1b365d", "#476a92", "#6e8db3", "#8ea8c6", "#b4c5d9", "#d8e1ec"]
    for index, (label, value) in enumerate(items):
        width = 520 * (value / max_value if max_value else 0)
        bars.append(f'<text x="70" y="{y + 18}" class="bar-label">{_escape_svg(label)}</text>')
        bars.append(f'<rect x="250" y="{y}" width="{width:.1f}" height="28" rx="8" fill="{palette[index % len(palette)]}"></rect>')
        bars.append(f'<text x="{260 + width:.1f}" y="{y + 18}" class="bar-value">{value} LOC</text>')
        y += 38
    return f'''<svg class="bars-svg" viewBox="0 0 900 320" role="img" aria-label="Distribuzione del perimetro core per linguaggio">
  <text x="450" y="32" text-anchor="middle" class="heat-title">Distribuzione del perimetro core</text>
  {''.join(bars)}
</svg>'''


def _report_context(context: dict[str, Any]) -> dict[str, Any]:
    settings = get_settings()
    summary = context.get("summary", {}) if isinstance(context.get("summary", {}), dict) else {}
    tool_results = context.get("tool_results", {}) if isinstance(context.get("tool_results", {}), dict) else {}
    reference_documents = context.get("reference_documents") if isinstance(context.get("reference_documents"), dict) else load_reference_bundle(settings.reference_root, str(context.get("client_name", "")), str(context.get("project_name", "")))
    llm_path = context.get("llm_output_path")
    if not llm_path and isinstance(summary.get("llm"), dict):
        llm_path = summary["llm"].get("output")
    llm_report_body = _sanitize_llm_markdown(_read_text(llm_path))
    effective_body = llm_report_body or _fallback_body(context)
    completed_tools, failed_tools, skipped_tools = _tool_status_counts(tool_results)
    enriched = dict(context)
    manifest_payload = enriched.get("manifest", {}) if isinstance(enriched.get("manifest", {}), dict) else {}
    analysis_path = manifest_payload.get("source") or manifest_payload.get("target") or enriched.get("project_name", "target")
    enriched.setdefault("generated_at", _italian_date())
    enriched.setdefault("report_title", TITLE)
    enriched.setdefault("classification", CLASSIFICATION)
    enriched.setdefault("report_outline", _outline())
    enriched.setdefault("security_view_standards", SECURITY_VIEW_STANDARDS)
    enriched.setdefault("compliance_standards", COMPLIANCE_STANDARDS)
    enriched.setdefault("scope_label", _scope_label(context))
    enriched.setdefault("summary_highlights", summary.get("highlights", []) if isinstance(summary.get("highlights", []), list) else [])
    enriched["completed_tools"] = completed_tools
    enriched["failed_tools"] = failed_tools
    enriched["skipped_tools"] = skipped_tools
    enriched["llm_report_body"] = effective_body
    enriched["report_sections"] = _extract_sections(effective_body)
    enriched["topology_svg"] = _topology_svg(enriched)
    enriched["heatmap_svg"] = _heatmap_svg(enriched)
    enriched["language_svg"] = _language_svg(enriched)
    enriched["reference_documents"] = reference_documents
    enriched["analysis_environment"] = _environment_summary(enriched)
    enriched["recipients_label"] = _recipient_label(enriched)
    enriched["overall_rating"] = _overall_rating(enriched)
    enriched["legend_rows"] = _legend_rows()
    enriched["rating_dimension_rows"] = _rating_dimension_rows()
    enriched["delivery_inventory_rows"] = _delivery_inventory_rows()
    enriched["rating_logic_rows"] = _rating_logic_rows()
    enriched["grade_meaning_rows"] = _grade_meaning_rows()
    enriched["rating_logic_intro"] = _rating_logic_intro()
    enriched["rating_logic_clauses"] = _rating_logic_clauses()
    enriched["disclaimer_text"] = _disclaimer_text()
    enriched["evidence_manifest_intro"] = _evidence_manifest_intro()
    enriched["evidence_manifest_json"] = json.dumps(manifest_payload, ensure_ascii=False, indent=2)
    enriched["finding_definition_block"] = _finding_definition_block()
    enriched["analysis_target_path"] = str(Path(str(analysis_path)))
    return enriched


def _relative_asset_path(asset_path: Path, base_dir: Path) -> str:
    return Path(os.path.relpath(asset_path, start=base_dir)).as_posix()


def _write_visual_assets(asset_dir: Path, context: dict[str, Any]) -> dict[str, Path]:
    assets: dict[str, Path] = {}
    try:
        import cairosvg
    except Exception:
        return assets
    payloads = {
        "topology": (context.get("topology_svg"), 1800, 900),
        "heatmap": (context.get("heatmap_svg"), 1800, 1400),
        "languages": (context.get("language_svg"), 1800, 640),
    }
    asset_dir.mkdir(parents=True, exist_ok=True)
    for stem, (svg_payload, width, height) in payloads.items():
        if not isinstance(svg_payload, str) or "<svg" not in svg_payload:
            continue
        target = asset_dir / f"{stem}.png"
        try:
            cairosvg.svg2png(
                bytestring=svg_payload.encode("utf-8"),
                write_to=str(target),
                output_width=width,
                output_height=height,
            )
        except Exception:
            continue
        assets[stem] = target
    return assets


def _write_visual_assets_matplotlib(asset_dir: Path, context: dict[str, Any]) -> dict[str, Path]:
    assets: dict[str, Path] = {}
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import matplotlib.patches as mpatches
    except Exception as exc:
        raise RuntimeError("matplotlib is required to render report visual assets") from exc
    asset_dir.mkdir(parents=True, exist_ok=True)
    blue = f"#{PRIMARY_BLUE}"
    sec_blue = f"#{SECONDARY_BLUE}"
    codebase_metrics = context.get("codebase_metrics", {}) if isinstance(context.get("codebase_metrics", {}), dict) else {}
    repo_profile = context.get("repo_profile", {}) if isinstance(context.get("repo_profile", {}), dict) else {}
    risk_register = context.get("risk_register_payload", []) if isinstance(context.get("risk_register_payload"), list) else []
    if not risk_register:
        risk_register = context.get("executive_vulnerabilities", []) if isinstance(context.get("executive_vulnerabilities"), list) else []

    try:
        languages = repo_profile.get("languages", []) or []
        frameworks = repo_profile.get("frameworks", []) or []
        source_paths = repo_profile.get("source_paths", []) or ["target"]
        fig, ax = plt.subplots(figsize=(9, 4.2))
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 5)
        ax.axis("off")
        ax.add_patch(mpatches.FancyBboxPatch((3.5, 1.5), 3, 2, boxstyle="round,pad=0.2", facecolor=blue, edgecolor="none"))
        ax.text(5, 2.8, str(context.get("project_name", "Target")), ha="center", va="center", fontsize=12, color="white", fontweight="bold")
        ax.text(5, 2.2, ", ".join(str(Path(p).name) for p in source_paths[:3]), ha="center", va="center", fontsize=8, color="#dbe5f4")
        nodes = [("Linguaggi", ", ".join(languages[:3]), 8.2, 4), ("Framework", ", ".join(frameworks[:3]), 8.2, 2.5), ("Evidence", f"{codebase_metrics.get('files', 0)} file", 8.2, 1)]
        for label, value, nx, ny in nodes:
            ax.add_patch(mpatches.FancyBboxPatch((nx - 0.9, ny - 0.4), 1.8, 0.8, boxstyle="round,pad=0.1", facecolor="#eef4fb", edgecolor="#b6c7dd"))
            ax.text(nx, ny + 0.1, label, ha="center", va="center", fontsize=7, fontweight="bold", color="#354052")
            ax.text(nx, ny - 0.15, value[:28], ha="center", va="center", fontsize=6, color="#5a677f")
            ax.annotate("", xy=(nx - 0.9, ny), xytext=(6.5, 2.5), arrowprops=dict(arrowstyle="->", color="#8a97aa", lw=1))
        ax.add_patch(mpatches.FancyBboxPatch((0.5, 1.8), 2, 1.2, boxstyle="round,pad=0.1", facecolor="#eef4fb", edgecolor="#b6c7dd"))
        ax.text(1.5, 2.4, str(context.get("client_name", "Client")), ha="center", va="center", fontsize=9, color="#354052")
        ax.annotate("", xy=(3.5, 2.5), xytext=(2.5, 2.5), arrowprops=dict(arrowstyle="->", color="#5d6b80", lw=1.5))
        fig.tight_layout()
        path = asset_dir / "topology.png"
        fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        assets["topology"] = path
    except Exception:
        pass

    if risk_register:
        try:
            fig, ax = plt.subplots(figsize=(9, 7))
            ax.set_xlim(0.5, 5.5)
            ax.set_ylim(0.5, 5.5)
            ax.set_xlabel("Likelihood", fontsize=10, color="#4f5c72")
            ax.set_ylabel("Technical Impact", fontsize=10, color="#4f5c72")
            ax.set_title("Heatmap Rischi per Componenti", fontsize=14, fontweight="bold", color="#344055", pad=12)
            for x in range(1, 6):
                for y in range(1, 6):
                    alpha = 0.04 + 0.03 * (x + y - 2) / 8
                    ax.add_patch(plt.Rectangle((x - 0.45, y - 0.45), 0.9, 0.9, facecolor=blue, alpha=alpha, edgecolor="#d6dee8", lw=0.5))
            ax.axhline(y=3, color="#c0cede", lw=0.8, ls="--")
            ax.axvline(x=3, color="#c0cede", lw=0.8, ls="--")
            for finding in risk_register[:10]:
                lk = int(finding.get("likelihood", 3))
                ti = int(finding.get("technical_impact", 3))
                grade = str(finding.get("grade", "C"))
                color = "#c0392b" if grade in ("D", "E") else "#e67e22" if grade == "C" else sec_blue
                ax.scatter(lk, ti, s=280, color=color, edgecolors="white", linewidths=1.5, zorder=5)
                fid = str(finding.get("finding_id", ""))
                ax.annotate(fid, (lk, ti), textcoords="offset points", xytext=(8, 8), fontsize=6.5, color="#2f3948", fontweight="bold")
            ax.set_xticks(range(1, 6))
            ax.set_yticks(range(1, 6))
            ax.tick_params(labelsize=8)
            fig.tight_layout()
            path = asset_dir / "heatmap.png"
            fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
            plt.close(fig)
            assets["heatmap"] = path
        except Exception:
            pass

    by_language = codebase_metrics.get("by_language", {}) or {}
    if isinstance(by_language, dict) and by_language:
        try:
            items = sorted(((lang, counts.get("code_lines", 0)) for lang, counts in by_language.items() if isinstance(counts, dict)), key=lambda i: i[1], reverse=True)[:6]
            if items:
                labels, values = zip(*items)
                fig, ax = plt.subplots(figsize=(9, 3.2))
                palette = [blue, "#476a92", "#6e8db3", "#8ea8c6", "#b4c5d9", "#d8e1ec"]
                ax.barh(list(reversed(labels)), list(reversed(values)), color=[palette[i % len(palette)] for i in range(len(items) - 1, -1, -1)], height=0.6)
                ax.set_xlabel("LOC", fontsize=9)
                ax.set_title("Distribuzione del Perimetro Core", fontsize=13, fontweight="bold", color="#344055", pad=10)
                ax.tick_params(labelsize=8)
                for spine in ("top", "right"):
                    ax.spines[spine].set_visible(False)
                fig.tight_layout()
                path = asset_dir / "languages.png"
                fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
                plt.close(fig)
                assets["languages"] = path
        except Exception:
            pass
    return assets


def _insert_after_heading(markdown_text: str, heading_number: str, block: str) -> str:
    if not block.strip():
        return markdown_text
    lines = markdown_text.splitlines()
    marker = f"## {heading_number}."
    for index, line in enumerate(lines):
        if line.startswith(marker):
            lines[index:index + 1] = [line, "", block.strip(), ""]
            return "\n".join(lines).strip() + "\n"
    return markdown_text


def _inject_visual_blocks(markdown_text: str, assets: dict[str, Path], base_dir: Path) -> str:
    updated = markdown_text
    if assets.get("topology") or assets.get("heatmap"):
        blocks: list[str] = []
        if assets.get("topology"):
            blocks.append(f'![Fig. 3.1 - Topologia del perimetro applicativo]({_relative_asset_path(assets["topology"], base_dir)})')
        if assets.get("heatmap"):
            blocks.append(f'![Fig. 3.2 - Heatmap rischi per componenti]({_relative_asset_path(assets["heatmap"], base_dir)})')
        updated = _insert_after_heading(updated, "3", "\n\n".join(blocks))
    if assets.get("languages"):
        updated = _insert_after_heading(
            updated,
            "9",
            f'![Fig. 9.1 - Distribuzione del perimetro core]({_relative_asset_path(assets["languages"], base_dir)})',
        )
    return updated


def _inject_structural_blocks(markdown_text: str, context: dict[str, Any]) -> str:
    updated = markdown_text
    finding_block = str(context.get("finding_definition_block", "")).strip()
    if finding_block and "Definizione del finding" not in updated:
        updated = _insert_after_heading(updated, "11", finding_block)
    return updated


def render_report(template_root: Path, context: dict[str, object], output_path: Path) -> Path:
    env = Environment(loader=FileSystemLoader(template_root), autoescape=select_autoescape(enabled_extensions=("j2",)), trim_blocks=True, lstrip_blocks=True)
    env.filters["slug_anchor"] = _slug_anchor
    template = env.get_template("ai_technical_intelligence_review.md.j2")
    enriched = _report_context(context)
    raw_markdown = template.render(**enriched).strip() + "\n"
    asset_dir = output_path.parent / "_report_assets"
    assets = _write_visual_assets(asset_dir, enriched)
    if not assets:
        assets = _write_visual_assets_matplotlib(asset_dir, enriched)
    final_markdown = _inject_visual_blocks(raw_markdown, assets, output_path.parent)
    final_markdown = _inject_structural_blocks(final_markdown, enriched)
    output_path.write_text(final_markdown, encoding="utf-8")
    return output_path


def _reference_docx_path(context: dict[str, object]) -> Path:
    reference_documents = context.get("reference_documents", {}) if isinstance(context.get("reference_documents", {}), dict) else {}
    candidate_value = str(reference_documents.get("template_docx_path", "")).strip()
    if candidate_value:
        candidate = Path(candidate_value).expanduser()
        if candidate.exists():
            return candidate
    settings = get_settings()
    fallback = settings.reference_root / "template.docx"
    if fallback.exists():
        return fallback
    raise FileNotFoundError("template.docx not found in reference materials")


def _render_docx_via_pandoc(report_markdown_path: Path, context: dict[str, object], output_path: Path) -> Path:
    pandoc_bin = shutil.which("pandoc")
    if not pandoc_bin:
        raise RuntimeError("pandoc is required to render the DOCX deliverable")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    reference_docx = _reference_docx_path(context)
    command = [
        pandoc_bin,
        str(report_markdown_path.name),
        "--from",
        "gfm+pipe_tables+yaml_metadata_block",
        "--standalone",
        "--wrap=none",
        "--reference-doc",
        str(reference_docx),
        "--resource-path",
        str(report_markdown_path.parent),
        "--output",
        str(output_path),
    ]
    completed = subprocess.run(
        command,
        cwd=str(report_markdown_path.parent),
        check=False,
        capture_output=True,
        text=True,
    )
    if completed.returncode != 0:
        raise RuntimeError((completed.stderr or completed.stdout).strip() or "pandoc failed to render DOCX")
    return output_path


def _body_markdown_from_report(report_markdown_path: Path) -> str:
    lines = report_markdown_path.read_text(encoding="utf-8").splitlines()
    for index, line in enumerate(lines):
        if line.startswith("## 1. "):
            return "\n".join(lines[index:]).strip() + "\n"
    return report_markdown_path.read_text(encoding="utf-8")


def _strip_markdown_inline(text: str) -> str:
    cleaned = html.unescape(text.strip())
    cleaned = re.sub(r"!\[(.*?)\]\((.*?)\)", r"\1", cleaned)
    cleaned = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", cleaned)
    cleaned = re.sub(r"`([^`]*)`", r"\1", cleaned)
    cleaned = re.sub(r"\*\*([^*]+)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"\*([^*]+)\*", r"\1", cleaned)
    cleaned = re.sub(r"__([^_]+)__", r"\1", cleaned)
    cleaned = re.sub(r"_([^_]+)_", r"\1", cleaned)
    return re.sub(r"\s+", " ", cleaned).strip()


def _iter_template_paragraphs(document) -> list[object]:
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def _replace_paragraph_text(paragraph, value: str) -> None:
    from docx.oxml.ns import qn

    node = paragraph._p
    for child in list(node):
        if child.tag != qn("w:pPr"):
            node.remove(child)
    if value:
        paragraph.add_run(value)


def _set_paragraph_if_matches(paragraph, replacements: dict[str, str]) -> None:
    text = paragraph.text or ""
    updated = text
    for source, target in replacements.items():
        if source in updated:
            updated = updated.replace(source, target)
    if updated != text:
        _replace_paragraph_text(paragraph, updated)


def _front_matter_replacements(context: dict[str, object]) -> dict[str, str]:
    return {
        "[Titolo del Report / Documento]": TITLE,
        "[Nome Cliente / Progetto]": str(context.get("client_name") or context.get("project_name") or "Target"),
        "[Data di emissione]": str(context.get("generated_at") or _italian_date()),
    }


def _trim_template_body(document) -> None:
    from docx.oxml.ns import qn

    body = document._element.body
    preserved: list[object] = []
    non_empty_paragraphs = 0
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            preserved.append(child)
            continue
        if child.tag == qn("w:p"):
            text = "".join(node.text or "" for node in child.iter(qn("w:t"))).strip()
            if text:
                non_empty_paragraphs += 1
        if non_empty_paragraphs <= 3:
            preserved.append(child)
    for child in list(body):
        if child not in preserved and child.tag != qn("w:sectPr"):
            body.remove(child)


def _set_run_font(run, font_name: str, size_pt: float, *, bold: bool = False, color_hex: str = "2F3948", italic: bool = False) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color_hex)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), font_name)


def _set_paragraph_shading(paragraph, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _apply_paragraph_appearance(paragraph, kind: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm, Pt

    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1.15
    fmt.keep_together = True
    if kind == "cover_title":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.space_after = Pt(4)
    elif kind == "cover_meta":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.space_after = Pt(2)
    elif kind in {"title", "heading1"}:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.space_before = Pt(8)
        fmt.space_after = Pt(8)
        fmt.keep_with_next = True
    elif kind == "heading2":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.space_before = Pt(6)
        fmt.space_after = Pt(4)
        fmt.keep_with_next = True
    elif kind == "toc":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.space_after = Pt(2)
    elif kind == "toc_sub":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.left_indent = Mm(6)
        fmt.space_after = Pt(1)
    elif kind == "caption":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.space_before = Pt(2)
        fmt.space_after = Pt(6)
    elif kind == "code":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.left_indent = Mm(4)
        fmt.right_indent = Mm(4)
        fmt.space_after = Pt(0)
        _set_paragraph_shading(paragraph, "F4F7FB")
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _add_section_rule(document) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), PRIMARY_BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_rich_text_paragraph(container, text: str, kind: str = "body"):
    from docx.shared import Pt

    paragraph = container.add_paragraph()
    _apply_paragraph_appearance(paragraph, kind)
    if kind in {"code", "caption", "cover_title", "cover_meta", "toc", "toc_sub"}:
        run = paragraph.add_run(_strip_markdown_inline(text))
        _set_run_font(run, CODE_FONT if kind == "code" else BODY_FONT, 10.5 if kind == "code" else 11, color_hex="374151" if kind == "code" else "2F3948")
        return paragraph
    segments = re.split(r"(\*\*[^*]+\*\*|`[^`]+`|>[^\n]+)", text)
    for segment in segments:
        if not segment:
            continue
        if segment.startswith("**") and segment.endswith("**"):
            run = paragraph.add_run(segment[2:-2])
            _set_run_font(run, BODY_FONT, 11, bold=True, color_hex="2F3948")
        elif segment.startswith("`") and segment.endswith("`"):
            run = paragraph.add_run(segment[1:-1])
            _set_run_font(run, "Courier New", 10, color_hex="374151")
        elif segment.startswith(">"):
            run = paragraph.add_run(segment[1:].strip())
            _set_run_font(run, BODY_FONT, 11, italic=True, color_hex="526173")
        else:
            run = paragraph.add_run(_strip_markdown_inline(segment))
            _set_run_font(run, BODY_FONT, 11, color_hex="2F3948")
    return paragraph


def _clear_paragraph(paragraph) -> None:
    from docx.oxml.ns import qn

    node = paragraph._p
    for child in list(node):
        if child.tag != qn("w:pPr"):
            node.remove(child)


def _add_text_paragraph(container, text: str, kind: str = "body"):
    paragraph = container.add_paragraph()
    _apply_paragraph_appearance(paragraph, kind)
    run = paragraph.add_run(_strip_markdown_inline(text))
    if kind == "cover_title":
        _set_run_font(run, TITLE_FONT, 18, bold=True, color_hex=PRIMARY_BLUE)
    elif kind == "cover_meta":
        _set_run_font(run, BODY_FONT, 11, color_hex="526173")
    elif kind == "title":
        _set_run_font(run, HEADING_FONT, 16, bold=True, color_hex=PRIMARY_BLUE)
    elif kind == "heading1":
        _set_run_font(run, HEADING_FONT, 16, bold=True, color_hex=PRIMARY_BLUE)
    elif kind == "heading2":
        _set_run_font(run, SUBHEADING_FONT, 14, bold=True, color_hex=PRIMARY_BLUE)
    elif kind in {"toc", "toc_sub"}:
        _set_run_font(run, BODY_FONT, 11, color_hex=PRIMARY_BLUE)
    elif kind == "caption":
        _set_run_font(run, BODY_FONT, 10, italic=True, color_hex="526173")
    elif kind == "code":
        _set_run_font(run, CODE_FONT, 10.5, color_hex="374151")
    else:
        _set_run_font(run, BODY_FONT, 11, color_hex="2F3948")
    return paragraph


def _set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_border(cell, color: str = TABLE_BORDER) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        edge_tag = qn(f"w:{edge}")
        element = borders.find(edge_tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def _set_cell_text(cell, text: str, *, header: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    _clear_paragraph(paragraph)
    _apply_paragraph_appearance(paragraph, "body")
    run = paragraph.add_run(_strip_markdown_inline(text))
    if header:
        _set_run_font(run, HEADING_FONT, 11, bold=True, color_hex="FFFFFF")
        _set_cell_shading(cell, PRIMARY_BLUE)
    else:
        _set_run_font(run, BODY_FONT, 10.5, color_hex="2F3948")
    _set_cell_border(cell)


def _style_table(table) -> None:
    table.autofit = True
    with contextlib.suppress(Exception):
        table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            _set_cell_border(cell)
        if row_index == 0:
            for cell in row.cells:
                _set_cell_shading(cell, PRIMARY_BLUE)


def _add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _append_page_field(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run("Pagina ")
    _set_run_font(run, BODY_FONT, 9.5, color_hex="526173")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    field_run = paragraph.add_run()
    field_run._r.append(fld_begin)
    field_run._r.append(instr)
    field_run._r.append(fld_end)


def _configure_page_numbering(document) -> None:
    for section in document.sections:
        section.different_first_page_header_footer = True
        footer = section.footer
        footer.is_linked_to_previous = False
        if not any(p.text.strip() for p in footer.paragraphs):
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            _append_page_field(paragraph)


def _style_cover_page(document) -> None:
    paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    if not paragraphs:
        return
    mapping = [("cover_title", 18), ("cover_meta", 14), ("cover_meta", 11)]
    for paragraph, (kind, size) in zip(paragraphs[:3], mapping, strict=False):
        text = paragraph.text
        _clear_paragraph(paragraph)
        _apply_paragraph_appearance(paragraph, kind)
        run = paragraph.add_run(text)
        if kind == "cover_title":
            _set_run_font(run, TITLE_FONT, size, bold=True, color_hex=PRIMARY_BLUE)
        else:
            _set_run_font(run, HEADING_FONT if size == 14 else BODY_FONT, size, bold=size == 14, color_hex="526173")


def _build_outline_entries(markdown_text: str) -> list[dict[str, object]]:
    entries: list[dict[str, object]] = []
    slug_counts: dict[str, int] = {}
    current_top = ""
    for raw_line in markdown_text.splitlines():
        stripped = raw_line.strip()
        heading = re.match(r"^(#{2,4})\s+(.*)$", stripped)
        if not heading:
            continue
        level = len(heading.group(1))
        title = _strip_markdown_inline(heading.group(2).strip())
        numeric = re.match(r"^(\d+)\.", title)
        if level == 2:
            current_top = numeric.group(1) if numeric else ""
        if level == 3 and current_top == "1":
            continue
        if level > 3:
            continue
        if level == 3 and not current_top:
            continue
        base = f"bm_{_slug_anchor(title)}"
        slug_counts[base] = slug_counts.get(base, 0) + 1
        anchor = base if slug_counts[base] == 1 else f"{base}_{slug_counts[base]}"
        entries.append({"title": title, "level": level, "anchor": anchor})
    return entries


def _add_internal_link(paragraph, text: str, anchor: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), PRIMARY_BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        r_fonts.set(qn(f"w:{attr}"), BODY_FONT)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(int(11 * 2)))
    r_pr.extend([r_fonts, color, underline, size])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(r_pr)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _metadata_rows(context: dict[str, object]) -> list[tuple[str, str]]:
    manifest = context.get("manifest", {}) if isinstance(context.get("manifest", {}), dict) else {}
    repo_profile = context.get("repo_profile", {}) if isinstance(context.get("repo_profile", {}), dict) else {}
    codebase_metrics = context.get("codebase_metrics", {}) if isinstance(context.get("codebase_metrics", {}), dict) else {}
    languages = ", ".join(repo_profile.get("languages", []) or []) or "Stack non determinato"
    return [
        ("Attività", "Digital Audit"),
        ("Repository di riferimento", str(manifest.get("source") or context.get("project_name") or "Target")),
        ("Hash SHA-512", str(manifest.get("repo_sha512") or "n/d")),
        ("Ambiente di analisi", str(context.get("analysis_environment") or "n/d")),
        ("Destinatari", str(context.get("recipients_label") or context.get("client_name") or "Cliente")),
        ("Tipo applicativo", _application_type(repo_profile)),
        ("Linguaggi core", languages),
        ("Righe di codice core", str(codebase_metrics.get("code_lines", 0))),
        ("Pipeline AI", "Normalizzazione, Clustering, Triage, Reporting"),
    ]


def _render_metadata_page(document, context: dict[str, object], outline_entries: list[dict[str, object]]) -> None:
    document.add_page_break()
    _add_text_paragraph(document, "METADATI DELL'ANALISI", "title")
    metadata = _metadata_rows(context)
    table = document.add_table(rows=1 + len(metadata), cols=2)
    _style_table(table)
    _set_cell_text(table.rows[0].cells[0], "Parametro", header=True)
    _set_cell_text(table.rows[0].cells[1], "Valore", header=True)
    for row_index, (label, value) in enumerate(metadata, start=1):
        cell_label = table.rows[row_index].cells[0]
        cell_label.text = ""
        paragraph = cell_label.paragraphs[0]
        _clear_paragraph(paragraph)
        _apply_paragraph_appearance(paragraph, "body")
        run = paragraph.add_run(_strip_markdown_inline(label))
        _set_run_font(run, BODY_FONT, 10.5, bold=True, color_hex="2F3948")
        _set_cell_border(cell_label)
        _set_cell_text(table.rows[row_index].cells[1], value)
        _set_cell_shading(cell_label, "E9F0F8")
    document.add_paragraph()
    _add_text_paragraph(document, "INDICE", "title")
    for entry in outline_entries:
        kind = "toc_sub" if int(entry["level"]) > 2 else "toc"
        paragraph = document.add_paragraph()
        _apply_paragraph_appearance(paragraph, kind)
        _add_internal_link(paragraph, str(entry["title"]), str(entry["anchor"]))
    document.add_page_break()


def _render_table_block(document, table_lines: list[str]) -> None:
    rows = []
    for raw_line in table_lines:
        stripped = raw_line.strip().strip("|")
        if not stripped:
            continue
        rows.append([_strip_markdown_inline(cell) for cell in stripped.split("|")])
    if len(rows) < 2:
        return
    headers = rows[0]
    data_rows = rows[2:] if len(rows) > 2 else []
    table = document.add_table(rows=1 + len(data_rows), cols=len(headers))
    _style_table(table)
    for index, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[index], header, header=True)
    for row_index, row in enumerate(data_rows, start=1):
        padded = row + [""] * (len(headers) - len(row))
        for column_index, value in enumerate(padded[: len(headers)]):
            _set_cell_text(table.rows[row_index].cells[column_index], value)


def _render_image_block(document, base_dir: Path, line: str) -> bool:
    match = re.match(r"!\[(.*?)\]\((.*?)\)", line.strip())
    if not match:
        return False
    alt_text, relative_path = match.groups()
    asset_path = (base_dir / relative_path).resolve()
    if not asset_path.exists():
        return False
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(asset_path), width=Mm(160))
    _add_text_paragraph(document, alt_text, "caption")
    return True


def _render_code_block(document, code_lines: list[str]) -> None:
    for line in code_lines or [""]:
        _add_text_paragraph(document, line or " ", "code")


def _render_markdown_into_docx(document, markdown_text: str, base_dir: Path, outline_entries: list[dict[str, object]]) -> None:
    lines = markdown_text.splitlines()
    outline_iter = iter(outline_entries)
    bookmark_id = 1
    index = 0
    first_top_level = True
    current_top = ""
    while index < len(lines):
        raw_line = lines[index].rstrip()
        stripped = raw_line.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("```"):
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index].rstrip("\n"))
                index += 1
            if index < len(lines):
                index += 1
            _render_code_block(document, code_lines)
            continue
        if stripped.startswith("|"):
            block = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                block.append(lines[index])
                index += 1
            _render_table_block(document, block)
            continue
        if _render_image_block(document, base_dir, stripped):
            index += 1
            continue
        heading = re.match(r"^(#{2,4})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = _strip_markdown_inline(heading.group(2).strip())
            numeric = re.match(r"^(\d+)\.", text)
            if level == 2:
                current_top = numeric.group(1) if numeric else ""
                if not first_top_level:
                    document.add_page_break()
                paragraph = _add_text_paragraph(document, text, "heading1")
                entry = next(outline_iter, None)
                if entry:
                    _add_bookmark(paragraph, str(entry["anchor"]), bookmark_id)
                    bookmark_id += 1
                _add_section_rule(document)
                first_top_level = False
            elif level == 3 and current_top == "1":
                _add_text_paragraph(document, text, "body")
            elif level == 3:
                paragraph = _add_text_paragraph(document, text, "heading2")
                entry = next(outline_iter, None)
                if entry:
                    _add_bookmark(paragraph, str(entry["anchor"]), bookmark_id)
                    bookmark_id += 1
            else:
                _add_text_paragraph(document, text, "body")
            index += 1
            continue
        if stripped.startswith("> "):
            _add_rich_text_paragraph(document, stripped, "body")
            index += 1
            continue
        if stripped == "---":
            _add_section_rule(document)
            index += 1
            continue
        if re.match(r"^-\s+", stripped) or re.match(r"^\d+\.\s+", stripped):
            while index < len(lines):
                current = lines[index].strip()
                if not re.match(r"^(-\s+|\d+\.\s+)", current):
                    break
                _add_rich_text_paragraph(document, current, "body")
                index += 1
            continue
        block = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].rstrip()
            candidate_stripped = candidate.strip()
            if not candidate_stripped:
                index += 1
                break
            if candidate_stripped.startswith("```") or candidate_stripped.startswith("|") or candidate_stripped.startswith("![") or re.match(r"^(#{2,4})\s+", candidate_stripped) or re.match(r"^-\s+", candidate_stripped) or re.match(r"^\d+\.\s+", candidate_stripped) or candidate_stripped == "---":
                break
            block.append(candidate_stripped)
            index += 1
        _add_rich_text_paragraph(document, " ".join(block), "body")


def render_docx(report_markdown_path: Path, context: dict[str, object], output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        from docx import Document
    except Exception:
        return _render_docx_via_pandoc(report_markdown_path, context, output_path)

    template_path = _reference_docx_path(context)
    document = Document(str(template_path))
    for paragraph in _iter_template_paragraphs(document):
        _set_paragraph_if_matches(paragraph, _front_matter_replacements(context))
    _trim_template_body(document)
    _style_cover_page(document)
    _configure_page_numbering(document)
    body_markdown = _body_markdown_from_report(report_markdown_path)
    outline_entries = _build_outline_entries(body_markdown)
    _render_metadata_page(document, context, outline_entries)
    _render_markdown_into_docx(document, body_markdown, report_markdown_path.parent, outline_entries)
    document.save(str(output_path))
    return output_path


def _fallback_pdf(report_markdown_path: Path, context: dict[str, object], output_path: Path) -> Path:
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleStyle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=18, leading=24, spaceAfter=10)
    body_style = ParagraphStyle("BodyStyle", parent=styles["BodyText"], fontName="Helvetica", fontSize=10, leading=14, spaceAfter=6)
    doc = SimpleDocTemplate(str(output_path), pagesize=A4, leftMargin=16 * mm, rightMargin=16 * mm, topMargin=18 * mm, bottomMargin=16 * mm, title=TITLE, author="Tecnolife")
    content = report_markdown_path.read_text(encoding="utf-8")
    story = [Paragraph(TITLE, title_style), Paragraph(html.escape(str(context.get("client_name", "Target Company"))), body_style), Spacer(1, 4 * mm)]
    for paragraph in [chunk.strip() for chunk in content.split("\n\n") if chunk.strip()]:
        safe = html.escape(paragraph).replace("\n", "<br/>")
        story.append(Paragraph(safe, body_style))
    doc.build(story)
    return output_path


def render_pdf(report_docx_path: Path, context: dict[str, object], output_path: Path, markdown_fallback_path: Path | None = None) -> Path:
    office_bin = shutil.which("libreoffice") or shutil.which("soffice")
    if not office_bin:
        if markdown_fallback_path is None:
            raise RuntimeError("LibreOffice is required to render PDF and no markdown fallback was provided")
        return _fallback_pdf(markdown_fallback_path, context, output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="niobe-lo-") as tempdir:
        env = os.environ.copy()
        env.setdefault("HOME", tempdir)
        env.setdefault("XDG_RUNTIME_DIR", tempdir)
        fc_cache_bin = shutil.which("fc-cache")
        if fc_cache_bin:
            subprocess.run([fc_cache_bin, "-f"], check=False, capture_output=True, text=True, env=env)
        command = [
            office_bin,
            "--headless",
            "--convert-to",
            "pdf:writer_pdf_Export",
            "--outdir",
            tempdir,
            str(report_docx_path),
        ]
        completed = subprocess.run(command, check=False, capture_output=True, text=True, env=env)
        generated_pdf = Path(tempdir) / f"{report_docx_path.stem}.pdf"
        if completed.returncode != 0 or not generated_pdf.exists():
            if markdown_fallback_path is None:
                raise RuntimeError((completed.stderr or completed.stdout).strip() or "LibreOffice failed to render PDF")
            return _fallback_pdf(markdown_fallback_path, context, output_path)
        shutil.copy2(generated_pdf, output_path)
    return output_path



def _supporting_docx_context(title: str) -> dict[str, object]:
    return {
        "client_name": title,
        "project_name": title,
        "generated_at": _italian_date(),
        "reference_documents": {},
    }


def _supporting_markdown_payload(title: str, markdown_text: str) -> str:
    cleaned = markdown_text.strip()
    if cleaned.startswith("---"):
        return cleaned
    return f"---\ntitle: {title}\n---\n\n## {title}\n\n{cleaned}\n"


def _render_supporting_docx(title: str, markdown_text: str, output_path: Path) -> Path:
    from docx import Document

    context = _supporting_docx_context(title)
    template_path = _reference_docx_path(context)
    document = Document(str(template_path))
    for paragraph in _iter_template_paragraphs(document):
        _set_paragraph_if_matches(
            paragraph,
            {
                "[Titolo del Report / Documento]": title,
                "[Nome Cliente / Progetto]": title,
                "[Data di emissione]": _italian_date(),
            },
        )
    _trim_template_body(document)
    _style_cover_page(document)
    _configure_page_numbering(document)

    anchor = _slug_anchor(title)
    _render_metadata_page(document, context, [{"title": title, "level": 2, "anchor": anchor}])
    paragraph = _add_text_paragraph(document, title, "heading1")
    _add_bookmark(paragraph, anchor, 1)
    _add_section_rule(document)
    _render_markdown_into_docx(document, markdown_text.strip(), output_path.parent, [])
    document.save(str(output_path))
    return output_path


def render_supporting_pdf(title: str, markdown_text: str, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    temp_docx = output_path.with_suffix(".docx")
    temp_md = output_path.with_suffix(".tmp.md")
    temp_md.write_text(_supporting_markdown_payload(title, markdown_text), encoding="utf-8")
    try:
        try:
            _render_supporting_docx(title, markdown_text, temp_docx)
            return render_pdf(temp_docx, _supporting_docx_context(title), output_path, markdown_fallback_path=temp_md)
        except Exception:
            html_path = output_path.with_suffix(".tmp.html")
            html_payload = f"""<!DOCTYPE html>
<html lang=\"it\">
<head>
<meta charset=\"utf-8\" />
<style>
  @page {{ size: A4; margin: 18mm 15mm 18mm 15mm; }}
  body {{ font-family: Aptos, \"Liberation Sans\", \"DejaVu Sans\", sans-serif; color: #2f3948; font-size: 11pt; text-align: justify; }}
  h1 {{ color: #1b365d; font-size: 18pt; border-left: 4px solid #1b365d; padding-left: 12px; }}
  h2 {{ color: #1b365d; font-size: 16pt; margin-top: 18px; }}
  h3 {{ color: #1b365d; font-size: 14pt; margin-top: 14px; }}
  p, li {{ line-height: 1.55; text-align: justify; }}
  pre, code {{ font-family: \"Aptos Mono\", \"Courier New\", monospace; }}
  table {{ width: 100%; border-collapse: collapse; margin-top: 12px; }}
  th, td {{ border: 1px solid #a9b8c8; padding: 6px; vertical-align: top; }}
  th {{ background: #1b365d; color: #fff; text-align: left; }}
</style>
</head>
<body>
<h1>{html.escape(title)}</h1>
{_markdown_to_html(markdown_text)}
</body>
</html>"""
            html_path.write_text(html_payload, encoding="utf-8")
            try:
                from weasyprint import HTML
                HTML(filename=str(html_path), base_url=str(html_path.parent)).write_pdf(str(output_path))
            except Exception:
                styles = getSampleStyleSheet()
                doc = SimpleDocTemplate(str(output_path), pagesize=A4, leftMargin=16 * mm, rightMargin=16 * mm, topMargin=18 * mm, bottomMargin=16 * mm, title=title, author="Tecnolife")
                body_style = ParagraphStyle("BodyStyle", parent=styles["BodyText"], fontName="Helvetica", fontSize=10, leading=14, spaceAfter=6)
                story = [Paragraph(html.escape(title), ParagraphStyle("Title", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=18, leading=22, textColor=colors.HexColor("#1b365d"))), Spacer(1, 4 * mm)]
                for paragraph in [chunk.strip() for chunk in markdown_text.split("\n\n") if chunk.strip()]:
                    safe = html.escape(paragraph).replace("\n", "<br/>")
                    story.append(Paragraph(safe, body_style))
                doc.build(story)
            finally:
                with contextlib.suppress(FileNotFoundError):
                    html_path.unlink()
            return output_path
    finally:
        with contextlib.suppress(FileNotFoundError):
            temp_md.unlink()
        with contextlib.suppress(FileNotFoundError):
            temp_docx.unlink()
